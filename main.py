# -----------------------------------------------------------------------------------------------------------------
# AllToPDF
# 一站式多格式文件转PDF转换器

# 详细介绍：
# AllToPDF是一个强大的Python脚本，用于将多种文件格式（如JPG, PNG, DOCX, TXT等）转换为PDF格式。
# 它能自动识别文件类型并进行相应的转换处理，适合需要将大量不同类型文件统一为PDF格式的场景，如办公室文档管理、
# 学术资料整理等。

# 使用方法：
# 1. 环境配置：确保Python环境已安装，并安装所需的库，包括pymupdf、Pillow、docx2pdf等。
# 2. 路径设置：在脚本中指定输入文件夹的路径（file_path）和输出PDF文件的路径（output_pdf_path）。
# 3. 运行脚本：执行脚本后，它会自动处理指定路径下的所有文件，将它们转换为PDF，并保存在输出路径。

# 注意：
# - 本程序不支持Excel格式（xls，xlsx)
# - 文件夹不能为空
# - 文件支持：确保输入文件是脚本支持的格式。当前支持的格式包括图片（JPG, PNG, GIF）、文档（DOC, DOCX, TXT）和PDF。
# - 错误处理：脚本将跳过无法处理的文件，并在最后提供错误计数。
# - 性能因素：处理大量或大尺寸文件时可能需要较长时间。

# © Yuping Pan 2023-2024
# -----------------------------------------------------------------------------------------------------------------

import fitz  # 注：导入本模块需安装pymupdf库
import os
from natsort import natsorted
from PIL import Image
from docx2pdf import convert
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Color print
from termcolor import colored, cprint


def get_file_name_listdir(file_dir):
    # 返回指定目录下的文件列表（包括文件、文件夹）
    return os.listdir(file_dir)


def is_folder(enter_path, enter_file_name):
    # 检查给定路径是否为文件夹
    path = enter_path + "/" + enter_file_name
    if os.path.isdir(path):
        return True
    else:
        return False

def get_file_extension(_file_path):
    """
    获取给定文件的文件扩展名。

    返回:
    str: 文件扩展名。
    """
    in_file_kind = os.path.splitext(_file_path)[1][1:].lower()

    return in_file_kind

def get_file_name(_file_path):
    """
    获取给定文件的文件名（不包括扩展名）。

    返回:
    str: 文件名。
    """
    return os.path.splitext(_file_path)[0]


# 处理TXT文件
def create_word_file(txt_file,_temp_path):
    # 从TXT文件创建Word文档
    doc = Document()

    template = Document('template.docx')
    for paragraph in template.paragraphs:
        if "{content}" in paragraph.text:  # 替换文本
            new_paragraph = doc.add_paragraph()
            new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            full_content = txt_file

            run = new_paragraph.add_run(paragraph.text.replace("{content}", full_content))
            run.font.size = Pt(22)
            run.font.name = 'Times New Romans'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.bold = False

            # Insert empty paragraphs before the new paragraph
            for _ in range(5):
                empty_paragraph = new_paragraph.insert_paragraph_before("")

    doc.save(_temp_path)

def generate_prefix(level, is_last):
    """
    根据层级和是否为最后一个元素生成文件树的前缀
    """
    prefix = ""
    if level > 0:
        prefix += "│   " * (level - 1)
        prefix += "├─ " if not is_last else "└─ "
    return prefix


def convert_folder_to_pdf(path, level=0, is_last=False):
    global error_count
    global total_count

    # 打印当前处理的文件夹名称
    if level == 1:
        print(path.split('/')[-1] + "/")

    # 将文件夹中的内容转换为PDF
    in_folder = get_file_name_listdir(path)
    in_folder = natsorted(in_folder)

    forbid_list = [".DS_Store", "Thumbs.db"]
    in_folder = [f for f in in_folder if f not in forbid_list]

    doc = fitz.open()

    for index, in_file in enumerate(in_folder):
        in_file_path = os.path.join(path, in_file)
        is_last_item = (index == len(in_folder) - 1)

        prefix = generate_prefix(level, is_last_item)

        total_count = total_count + 1  # 成功计数

        if is_folder(path, in_file):
            # 如果是文件夹，则递归处理
            print(prefix + in_file + "/")
            initial_pdf = convert_folder_to_pdf(in_file_path, level + 1, is_last_item)
            doc.insert_pdf(initial_pdf)

        else:
            in_file_kind = get_file_extension(path + '/' + in_file)
            in_file_name = get_file_name(in_file)
            # print('in_file_kind is:',in_file_kind)

            # 根据文件类型处理文件
            try:
                # 图片文件处理
                if in_file_kind in ["jpg", "jpeg", "png", "gif"]:
                    try:
                        im = Image.open(in_file_path)
                        # 检查图片尺寸，如果太大，则适当缩小
                        max_size = 8000  # 最大宽度或高度
                        if im.size[0] > max_size or im.size[1] > max_size:
                            im.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
                        im = im.rotate(0, expand=1)
                        im_converted = im.convert('RGB') if im.mode != 'RGB' else im
                        im_converted.save(in_file_path, quality=95, optimize=True)

                        imgdoc = fitz.open(in_file_path)
                        pdfbytes = imgdoc.convert_to_pdf()
                        imgpdf = fitz.open("pdf", pdfbytes)
                        doc.insert_pdf(imgpdf)
                    except Exception as e:
                        _print_text = f"ERR Loading {in_file}: {e}"
                        error_count = error_count + 1  # 错误计数
                        print(colored(_print_text, 'red'))

                # PDF文件处理
                elif in_file_kind == "pdf" :
                    try:
                        stream = bytearray(open(in_file_path, "rb").read())
                        pdf_fitz = fitz.open("pdf", stream)
                        doc.insert_pdf(pdf_fitz)
                        # print("Success:" + in_file)
                    except:
                        _print_text = "ERR Loading:" + in_file
                        error_count = error_count + 1  # 错误计数
                        print(colored(_print_text, 'red'))

                # Word文档处理
                elif in_file_kind == "doc" or in_file_kind == "docx" :
                    try:
                        temp_path = '2 Temp' + '/' + in_file +'.pdf'
                        convert(in_file_path, temp_path)

                        stream = bytearray(open(temp_path, "rb").read())
                        pdf_fitz = fitz.open("pdf", stream)
                        doc.insert_pdf(pdf_fitz)

                        # print("Success:" + in_file)
                    except:
                        _print_text = "ERR Loading:" + in_file
                        error_count = error_count + 1  # 错误计数
                        print(colored(_print_text, 'red'))

                # TXT文件处理
                elif in_file_kind == "txt" :
                    try:
                        temp_path_txt_to_word = '2 Temp' + '/' + in_file_name +'.docx'
                        create_word_file(in_file_name,temp_path_txt_to_word)

                        temp_path_word_to_pdf = '2 Temp' + '/' + in_file_name + '.pdf'
                        convert(temp_path_txt_to_word,temp_path_word_to_pdf)

                        stream = bytearray(open(temp_path_word_to_pdf, "rb").read())
                        pdf_fitz = fitz.open("pdf", stream)
                        doc.insert_pdf(pdf_fitz)

                        # print("Success:" + in_file)
                    except:
                        _print_text = "ERR Loading:" + in_file
                        error_count = error_count + 1  # 错误计数
                        print(colored(_print_text, 'red'))

                else:
                    _print_text ="ERR Unknown Type:" + in_file
                    error_count = error_count + 1  # 错误计数
                    print(colored(_print_text, 'red'))

                print(prefix + in_file + " [处理成功]")

            except AttributeError:
                _print_text = prefix + in_file + " [处理失败]"
                error_count = error_count + 1  # 错误计数
                print(colored(_print_text, 'red'))

    return doc

# 参数设置 ---------------------------------------------------------------------------

file_path = "/Users/panyp/PycharmProjects/AllToPDF/1 Folders"
output_pdf_path = "/Users/panyp/PycharmProjects/AllToPDF/3 Output"


# 主程序 ------------------------------------------------------------------------------

folder = get_file_name_listdir(file_path)
folder = natsorted(folder)

total_count = 0
error_count = 0

_forbid_list = [".DS_Store", "Thumbs.db"]

for __item in _forbid_list:
    try:
        folder.remove(__item)
    except ValueError:
        pass

for index, _folder in enumerate(folder):
    folder_file_path = os.path.join(file_path, _folder)
    is_last_folder = (index == len(folder) - 1)
    # print(f"处理文件夹: {_folder}")  # 打印当前正在处理的文件夹名称
    pdf = convert_folder_to_pdf(folder_file_path, 1, is_last_folder)

    # 保存PDF
    try:
        pdf.save(output_pdf_path + '/' + _folder + '.pdf')
    except ValueError:
        print(colored(f"ERR:{_folder}文件夹为空",'yellow'))
        error_count = error_count + 1


print(colored(f"共计处理（包括文件夹）:{total_count}",'white'))
print(colored(f"错误个数:{error_count}",'yellow'))